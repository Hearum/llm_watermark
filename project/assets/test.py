# Adjusting the table creation code to handle cases where rows have differing numbers of columns
from docx import Document
# Create a new Word Document
doc = Document()
doc.add_heading('大学生沉默螺旋现象调查结果', level=1)

# Define each table with headings and data
tables = {
    "1.1 基本信息统计": [
        ["项目", "选项", "统计人数", "百分比"],
        ["性别", "男", "29", "50.9%"],
        ["", "女", "28", "49.1%"],
        ["年级", "大一", "6", "10.5%"],
        ["", "大二", "18", "31.6%"],
        ["", "大三", "24", "42.1%"],
        ["", "大四", "9", "15.8%"],
        ["专业领域", "人文社科类", "16", "28.1%"],
        ["", "理工科类", "20", "35.1%"],
        ["", "商科类", "10", "17.5%"],
        ["", "艺术类", "6", "10.5%"],
        ["", "其他", "5", "8.8%"]
    ],
    "1.2 沉默螺旋现象体验": [
        ["你是否听说过“沉默螺旋”这一概念？"],
        ["选项", "统计人数", "百分比"],
        ["是", "35", "61.4%"],
        ["否", "22", "38.6%"],

        ["在公共议题讨论中，你是否曾经因为担心被孤立而选择保持沉默？"],
        ["选项", "统计人数", "百分比"],
        ["经常", "8", "14%"],
        ["有时", "22", "38.6%"],
        ["很少", "17", "29.8%"],
        ["从不", "10", "17.5%"],

        ["当你发现你的观点与大多数人不一致时，你通常会："],
        ["选项", "统计人数", "百分比"],
        ["保持沉默，避免表达自己的观点", "15", "26.3%"],
        ["尝试理解他人观点，但保持立场", "23", "40.4%"],
        ["积极表达自己的观点", "15", "26.3%"],
        ["其他（请说明）", "4", "7%"],

        ["你认为在大学校园中，哪些议题最容易出现沉默螺旋现象？"],
        ["选项", "统计人数", "百分比"],
        ["政治议题", "21", "36.8%"],
        ["社会议题", "16", "28.1%"],
        ["校园政策", "9", "15.8%"],
        ["学术争议", "7", "12.3%"],
        ["其他", "4", "7%"]
    ],
    "1.3 对沉默螺旋现象的看法": [
        ["你认为沉默螺旋现象对公共讨论的影响是："],
        ["选项", "统计人数", "百分比"],
        ["非常负面", "17", "29.8%"],
        ["有些负面", "24", "42.1%"],
        ["没有影响", "5", "8.8%"],
        ["有些正面", "8", "14%"],
        ["非常正面", "3", "5.3%"],

        ["你认为如何减少沉默螺旋现象的发生？"],
        ["选项", "统计人数", "百分比"],
        ["提高公众对沉默螺旋的认识", "14", "24.6%"],
        ["鼓励多元化观点的表达", "22", "38.6%"],
        ["建立更包容的讨论环境", "17", "29.8%"],
        ["其他", "4", "7%"]
    ],
    "具体的实验数据示例": [
        ["年级", "经常", "有时", "很少", "从不"],
        ["大一", "1", "2", "2", "1"],
        ["大二", "2", "6", "5", "5"],
        ["大三", "4", "10", "6", "4"],
        ["大四", "1", "4", "4", "0"]
    ]
}

# Add each table to the document
for section, rows in tables.items():
    doc.add_heading(section, level=2)
    
    # Create a table with columns equal to the maximum row length
    max_columns = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=max_columns)
    
    # Fill in the header row
    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(rows[0]):
        hdr_cells[i].text = heading
    
    # Fill in the rest of the rows
    for row in rows[1:]:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell
    doc.add_paragraph("")  # Add space between tables

# Save the documentoutput_path = r"F:\课程资料\24年秋大三上\综设\大学生沉默螺旋现象调查结果.docx"

output_path = "F:/课程资料/24年秋大三上/综设/大学生沉默螺旋现象调查结果.docx"
doc.save(output_path)

output_path
# import pandas as pd
# import ace_tools as tools

# # Creating a DataFrame to represent the table from the image
# data = {
#     "类别": ["时政新闻", "娱乐八卦", "文化教育", "商业经济", "生活健康", "其他"],
#     "个案数": [163, 163, 163, 163, 163, 163],
#     "范围": [6, 6, 5, 6, 6, 6],
#     "最小值": [0, 0, 0, 0, 0, 0],
#     "最大值": [6, 6, 5, 6, 6, 6],
#     "平均值": [1.79, 1.98, 1.98, 2.34, 2.45, 1.15],
#     "标准差": [1.312, 1.657, 1.217, 1.928, 2.091, 2.172],
#     "方差": [1.722, 2.747, 1.481, 3.719, 4.373, 4.719]
# }

# # Creating the DataFrame
# df = pd.DataFrame(data)

# # Displaying the table to the user
# tools.display_dataframe_to_user(name="社交平台关注话题情况", dataframe=df)

# import matplotlib.pyplot as plt
# import numpy as np

# # Data: score ranges and corresponding number of people
# score_ranges = ['0-5', '6-10', '11-15', '16-20', '21-25']

# # Hypothetical data: before and after opinion climate formation
# scores_no_climate = [5, 8, 12, 7, 3]  # Before opinion climate formation
# scores_with_climate = [1, 3, 5, 12, 16]  # After opinion climate formation

# # Set X-axis positions
# x = np.arange(len(score_ranges))

# # Set bar width
# width = 0.35

# # Create a bar chart with two sets of bars
# fig, ax = plt.subplots(figsize=(10, 6))
# bars1 = ax.bar(x - width/2, scores_no_climate, width, label='After Opinion Climate', color='skyblue')
# bars2 = ax.bar(x + width/2, scores_with_climate, width, label='Before Opinion Climate', color='salmon')

# # Set title and labels
# ax.set_title('Expression Willingness Distribution', fontsize=16)
# ax.set_xlabel('Score Range', fontsize=12)
# ax.set_ylabel('Number of People', fontsize=12)
# ax.set_xticks(x)
# ax.set_xticklabels(score_ranges)

# # Add value labels on top of each bar
# for bars in [bars1, bars2]:
#     for bar in bars:
#         height = bar.get_height()
#         ax.text(bar.get_x() + bar.get_width() / 2, height + 0.2, str(height), ha='center', fontsize=12)

# # Add legend
# ax.legend()

# # Show the plot
# plt.tight_layout()
# plt.show()
